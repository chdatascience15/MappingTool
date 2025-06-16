import streamlit as st
import pandas as pd
from sentence_transformers import SentenceTransformer
import numpy as np
from docx import Document
from docx.shared import RGBColor
import tempfile

import zipfile
import nltk

def ensure_punkt():
    punkt_dir = "nltk_data/tokenizers/punkt"
    zip_file = "punkt.zip"

    if not os.path.exists(punkt_dir):
        if os.path.exists(zip_file):
            st.info("Extracting punkt.zip...")
            with zipfile.ZipFile(zip_file, "r") as zip_ref:
                zip_ref.extractall("nltk_data")
        else:
            st.error("punkt.zip not found and punkt tokenizer is missing.")
            st.stop()

    nltk.data.path.append("nltk_data")

# Call early in your script before you use sent_tokenize()
ensure_punkt()

from nltk.tokenize import sent_tokenize
import os
os.environ["CUDA_VISIBLE_DEVICES"] = ""

# Ensure NLTK punkt tokenizer is available
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

# Load model
model = SentenceTransformer('all-MiniLM-L6-v2')


# Functions
def get_doc_name(uploaded_document_name):

    base_name = os.path.splitext(uploaded_document_name)[0]  # "QA Meeting Minutes 7-1-2024"

    output_name = base_name + ' (Mapped).docx'

    return output_name

# Find matching practices
def find_matches(text_chunk, threshold=0.5):
    chunk_embedding = model.encode([text_chunk], normalize_embeddings=True)[0]
    similarities = np.dot(practice_embeddings, chunk_embedding)
    matches = [(practice_names[i], float(similarities[i]))
            for i in range(len(similarities)) if similarities[i] >= threshold]
    return matches

# Process document text
def process_text(text):
    chunks = [" ".join(sent_tokenize(text)[i:i+2]) for i in range(0, len(sent_tokenize(text)), 2)]
    processed = []
    matched_practices_set = set()

    for chunk in chunks:
        matches = find_matches(chunk)
        if matches:
            match_names = [m[0] for m in matches]
            matched_practices_set.update(match_names)
            processed.append(f"{chunk} ({', '.join(match_names)})")
        else:
            processed.append(chunk)

    return processed, sorted(matched_practices_set)

# Generate Word document
def generate_doc(processed_chunks, matched_practices):
    doc = Document()

    # Add matched practice matrix
    doc.add_heading("Matched Groups", level=1)
    for p in matched_practices:
        doc.add_paragraph(p)

    doc.add_page_break()
    doc.add_heading("Document Text", level=1)

    for chunk in processed_chunks:
        para = doc.add_paragraph(chunk)
        if '(' in chunk:
            run = para.runs[0]
            run.font.color.rgb = RGBColor(0, 128, 0)  # Green for matched

    # Save to temp file
    tmp_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp_path.name)
    return tmp_path.name

# Streamlit UI
st.title("Text Matching Tool")


# Load Practice Statements

mapping_file = st.file_uploader("Upload your file with the Group Mappings", type=["csv"], help="Upload a CSV with a column with groupds and the corresponding statements to evaluate")
#mapping_file = st.text_input("Enter the path to your CSV file with group mappings:", placeholder="Enter Full path here")

st.write(mapping_file)
df = pd.DataFrame()

# if df.empty:
#     st.warning("Please upload a CSV file with group mappings and statements to proceed.")
# else:
if mapping_file:
    df = pd.read_csv(mapping_file)

group_column = st.selectbox("Select the column with the Group Mappings", options=df.columns.tolist(), help="Select the column that contains the group names or identifiers.")
statement_column = st.selectbox("Select the column with the Statements", options=df.columns.tolist(), help="Select the column that contains the practice statements or descriptions.")


# Assign columns for embedding and naming
practice_names = df[group_column].tolist()  # e.g., "Project Monitoring and Control"
practice_texts = df[statement_column].tolist()     # descriptive text for embedding

# Embed practice values
practice_embeddings = model.encode(practice_texts, normalize_embeddings=True)

# Streamlit input for document text
input_method = st.radio("Select input method:", ["Paste text", "Upload Word document"])

doc_text = ""
if input_method == "Paste text":
    doc_title = st.text_input("Document Title", placeholder="EX: Process and Product Quality Meeting Minutes") + " (Mapping).docx"
    doc_text = st.text_area("Paste your meeting transcript or notes below:")
elif input_method == "Upload Word document":
    uploaded_file = st.file_uploader("Upload a Word document (.docx)", type="docx")
    if uploaded_file:
        word_doc = Document(uploaded_file)
        doc_text = "\n".join([para.text for para in word_doc.paragraphs])
        doc_title = get_doc_name(uploaded_file.name)

if st.button("Analyze Document") and doc_text:
    st.write(f"Analyzing document: {doc_title}")
    with st.spinner("Processing document..."):
        processed_chunks, matched_practices = process_text(doc_text)
        output_path = generate_doc(processed_chunks, matched_practices)
        
        st.success("Analysis complete!")
        st.download_button(
            label="Download Annotated Word Document",
            data=open(output_path, "rb").read(),
            file_name=doc_title,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
